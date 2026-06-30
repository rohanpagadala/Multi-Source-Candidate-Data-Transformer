import unittest
import os
from src.extractors.resume import ResumeExtractor
from src.extractors.notes import NotesExtractor
from src.extractors.ats import ATSExtractor
from src.extractors.csv_export import CSVExtractor
from src.extractors.linkedin import LinkedInExtractor
from src.extractors.github import GitHubExtractor

class TestExtractors(unittest.TestCase):
    def setUp(self):
        self.data_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data")

    def test_resume_txt(self):
        path = os.path.join(self.data_dir, "resume.txt")
        extractor = ResumeExtractor()
        with open(path, "rb") as f:
            content = f.read()
        
        self.assertTrue(extractor.can_extract(path, content))
        facts = extractor.extract(path, content)
        
        # Verify basic facts extracted
        fields = [f.field for f in facts]
        self.assertIn("full_name", fields)
        self.assertIn("email", fields)
        self.assertIn("phone", fields)
        self.assertIn("skill", fields)
        self.assertIn("experience_item", fields)
        
        name_fact = next(f for f in facts if f.field == "full_name")
        self.assertEqual(name_fact.value.title(), "Rohan Pagadala")

    def test_recruiter_notes(self):
        path = os.path.join(self.data_dir, "recruiter_notes.txt")
        extractor = NotesExtractor()
        with open(path, "rb") as f:
            content = f.read()
            
        self.assertTrue(extractor.can_extract(path, content))
        facts = extractor.extract(path, content)
        
        fields = [f.field for f in facts]
        self.assertIn("full_name", fields)
        self.assertIn("email", fields)
        self.assertIn("skill", fields)
        
        # Notes specifically finds "Faiss" and "Langchain" from phrase extraction
        skills = [f.value for f in facts if f.field == "skill"]
        self.assertIn("Faiss", skills)
        self.assertIn("Langchain", skills)

    def test_ats_json(self):
        path = os.path.join(self.data_dir, "ats_candidate.json")
        extractor = ATSExtractor()
        with open(path, "rb") as f:
            content = f.read()
            
        self.assertTrue(extractor.can_extract(path, content))
        facts = extractor.extract(path, content)
        
        fields = [f.field for f in facts]
        self.assertIn("full_name", fields)
        self.assertIn("email", fields)
        self.assertIn("phone", fields)
        self.assertIn("experience_item", fields)
        self.assertIn("education_item", fields)

    def test_csv_export(self):
        path = os.path.join(self.data_dir, "recruiter_export.csv")
        extractor = CSVExtractor()
        with open(path, "rb") as f:
            content = f.read()
            
        self.assertTrue(extractor.can_extract(path, content))
        facts = extractor.extract(path, content)
        
        fields = [f.field for f in facts]
        self.assertIn("full_name", fields)
        self.assertIn("email", fields)
        self.assertIn("phone", fields)
        self.assertIn("experience_item", fields)

    def test_linkedin(self):
        path = os.path.join(self.data_dir, "linkedin_profile.json")
        extractor = LinkedInExtractor()
        with open(path, "rb") as f:
            content = f.read()
            
        self.assertTrue(extractor.can_extract(path, content))
        facts = extractor.extract(path, content)
        
        fields = [f.field for f in facts]
        self.assertIn("full_name", fields)
        self.assertIn("headline", fields)
        self.assertIn("linkedin_link", fields)
        self.assertIn("skill", fields)

    def test_github(self):
        path = os.path.join(self.data_dir, "github_profile.json")
        extractor = GitHubExtractor()
        with open(path, "rb") as f:
            content = f.read()
            
        self.assertTrue(extractor.can_extract(path, content))
        facts = extractor.extract(path, content)
        
        fields = [f.field for f in facts]
        self.assertIn("full_name", fields)
        self.assertIn("github_link", fields)
        self.assertIn("skill", fields)

    def test_resume_docx(self):
        import docx
        from io import BytesIO
        doc = docx.Document()
        doc.add_paragraph("Rohan Pagadala")
        doc.add_paragraph("Email: rohan.pagadala15@gmail.com")
        doc.add_paragraph("Phone: 9876543210")
        doc.add_paragraph("Skills: Python, ML")
        
        f_bytes = BytesIO()
        doc.save(f_bytes)
        content = f_bytes.getvalue()
        
        extractor = ResumeExtractor()
        path = "test_resume.docx"
        self.assertTrue(extractor.can_extract(path, content))
        facts = extractor.extract(path, content)
        
        fields = [f.field for f in facts]
        self.assertIn("full_name", fields)
        self.assertIn("email", fields)
        self.assertIn("phone", fields)
        
        name_fact = next(f for f in facts if f.field == "full_name")
        self.assertEqual(name_fact.value, "Rohan Pagadala")

if __name__ == "__main__":
    unittest.main()
